import os
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup


def extract_text_from_file(filepath):
    """Extract text content from PDF, DOCX, TXT, or HTML files"""
    
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    
    # Check if file is actually HTML regardless of extension
    if is_html_file(filepath):
        return extract_from_html(filepath)
    
    try:
        if ext == '.pdf':
            return extract_from_pdf(filepath)
        elif ext == '.docx':
            return extract_from_docx(filepath)
        elif ext == '.txt':
            return extract_from_txt(filepath)
        elif ext == '.html' or ext == '.htm':
            return extract_from_html(filepath)
        else:
            # Try as text if unknown
            return extract_from_txt(filepath)
    except Exception as e:
        print(f"Error extracting text from {filepath}: {str(e)}")
        # Return empty string instead of raising to allow partial processing of batches
        return ""


def is_html_file(filepath):
    """Check if file content looks like HTML"""
    try:
        with open(filepath, 'rb') as f:
            start = f.read(1024).strip()
            # Check for common HTML markers
            if b'<!DOCTYPE html' in start or b'<html' in start or b'<body' in start:
                return True
    except:
        pass
    return False


def extract_from_html(filepath):
    """Extract text from HTML file"""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
                
            text = soup.get_text()
            
            # Break into lines and remove leading/trailing space on each
            lines = (line.strip() for line in text.splitlines())
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            # Drop blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
    except Exception as e:
        print(f"Error reading HTML file: {e}")
        return ""


def extract_from_pdf(filepath):
    """Extract text from PDF file"""
    text = ""
    try:
        reader = PdfReader(filepath)
        
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                print(f"Error reading page in PDF: {e}")
                continue
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return ""
    
    return text.strip()


def extract_from_docx(filepath):
    """Extract text from DOCX file"""
    try:
        doc = Document(filepath)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""


def extract_from_txt(filepath):
    """Extract text from TXT file with encoding fallback"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"Error reading TXT file: {e}")
            return ""
            
    return ""
